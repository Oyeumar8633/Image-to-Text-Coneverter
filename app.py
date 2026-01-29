"""
Image-to-Word Converter - Streamlit Application
Main application file for the OCR-based document converter.
"""

import streamlit as st
from PIL import Image
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import ocr_engine


def set_run_bold_italic(run, is_bold: bool, is_italic: bool):
    """
    Set bold and italic formatting for a docx run.
    
    Args:
        run: docx Run object
        is_bold: Boolean for bold formatting
        is_italic: Boolean for italic formatting
    """
    run.bold = is_bold
    run.italic = is_italic


def get_alignment_constant(alignment: str) -> int:
    """
    Convert alignment string to docx alignment constant.
    
    Args:
        alignment: String alignment ('left', 'center', 'right')
        
    Returns:
        docx alignment constant
    """
    alignment_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT
    }
    return alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)


def create_docx_from_blocks(text_blocks: list) -> Document:
    """
    Create a Word document from text blocks with formatting.
    
    Args:
        text_blocks: List of TextBlock objects
        
    Returns:
        Document object ready to be saved
    """
    doc = Document()
    
    for block in text_blocks:
        # Create paragraph
        paragraph = doc.add_paragraph()
        
        # Set alignment
        paragraph.alignment = get_alignment_constant(block.alignment)
        
        # Add text with formatting
        run = paragraph.add_run(block.text)
        set_run_bold_italic(run, block.is_bold, block.is_italic)
        
        # Set font size (default 11pt)
        run.font.size = Pt(11)
        
        # Add space after paragraph for readability
        paragraph.space_after = Pt(6)
    
    return doc


def create_docx_simple(text: str) -> Document:
    """
    Create a simple Word document from plain text (fallback).
    
    Args:
        text: Plain text string
        
    Returns:
        Document object
    """
    doc = Document()
    
    # Split text into paragraphs
    paragraphs = text.split('\n\n')
    
    for para_text in paragraphs:
        if para_text.strip():
            paragraph = doc.add_paragraph(para_text.strip())
            paragraph.space_after = Pt(6)
    
    return doc


def main():
    """Main Streamlit application."""
    st.set_page_config(
        page_title="Image-to-Word Converter",
        page_icon="📄",
        layout="wide"
    )
    
    st.title("📄 Image-to-Word Converter")
    st.markdown("### Extract text from images and convert to formatted Word documents")
    
    st.sidebar.header("Settings")
    
    # OCR Engine Selection
    ocr_engines = ["Tesseract OCR", "EasyOCR (Better for Handwriting)", "PaddleOCR (Best Overall)"]
    if not ocr_engine.EASYOCR_AVAILABLE:
        ocr_engines = [e for e in ocr_engines if "EasyOCR" not in e]
        st.sidebar.warning("⚠️ EasyOCR not installed. Install with: pip install easyocr")
    if not ocr_engine.PADDLEOCR_AVAILABLE:
        ocr_engines = [e for e in ocr_engines if "PaddleOCR" not in e]
        st.sidebar.warning("⚠️ PaddleOCR not installed. Install with: pip install paddlepaddle paddleocr")
    
    selected_engine = st.sidebar.selectbox(
        "OCR Engine",
        options=ocr_engines,
        help="PaddleOCR: Best overall accuracy. EasyOCR: Good for handwriting. Tesseract: Fast for printed text"
    )
    use_easyocr = selected_engine == "EasyOCR (Better for Handwriting)" and ocr_engine.EASYOCR_AVAILABLE
    use_paddleocr = selected_engine == "PaddleOCR (Best Overall)" and ocr_engine.PADDLEOCR_AVAILABLE
    
    # Preprocessing - Deep learning OCRs often work better WITHOUT preprocessing
    if use_easyocr or use_paddleocr:
        use_preprocessing = st.sidebar.checkbox(
            "Enable Image Preprocessing", 
            value=False,
            help="⚠️ Deep learning OCRs usually work better WITHOUT preprocessing. Only enable if needed."
        )
        # Confidence threshold for deep learning OCRs
        engine_name = "EasyOCR" if use_easyocr else "PaddleOCR"
        min_confidence = st.sidebar.slider(
            f"Minimum Confidence ({engine_name})",
            min_value=0.0,
            max_value=1.0,
            value=0.5,
            step=0.05,
            help="Higher values = more accurate but may miss text. Lower values = more text but may include errors."
        )
    else:
        use_preprocessing = st.sidebar.checkbox(
            "Enable Image Preprocessing", 
            value=True,
            help="Apply image enhancement to improve OCR accuracy"
        )
        min_confidence = 0.5  # Not used for Tesseract
    
    if use_preprocessing:
        preprocessing_method = st.sidebar.selectbox(
            "Preprocessing Method",
            options=["adaptive", "otsu", "morphology", "enhanced"],
            index=0,
            help="Different preprocessing strategies for various image types"
        )
    else:
        preprocessing_method = "adaptive"
    
    use_formatting = st.sidebar.checkbox(
        "Detect Formatting (Bold/Italic/Alignment)",
        value=True,
        help="Attempt to detect and preserve text formatting"
    )
    
    # Tesseract-specific settings (only show if Tesseract is selected)
    if not use_easyocr and not use_paddleocr:
        try_multiple_modes = st.sidebar.checkbox(
            "Try Multiple OCR Modes",
            value=False,
            help="Try different OCR modes and use the best result (slower but more accurate)"
        )
        
        if not try_multiple_modes:
            psm_mode = st.sidebar.selectbox(
                "OCR Mode (PSM)",
                options=[3, 6, 11, 13],
                index=1,
                format_func=lambda x: f"PSM {x} - {['Auto', 'Uniform Block', 'Sparse Text', 'Single Line'][[3,6,11,13].index(x)]}",
                help="Page Segmentation Mode: 6=Uniform block (default), 3=Auto, 11=Sparse text, 13=Single line"
            )
        else:
            psm_mode = 6  # Will be overridden
    else:
        try_multiple_modes = False
        psm_mode = 6
    
    st.sidebar.markdown("---")
    if use_paddleocr:
        st.sidebar.info("✅ **PaddleOCR** selected - Best overall accuracy for both printed and handwritten text")
    elif use_easyocr:
        st.sidebar.info("✅ **EasyOCR** selected - Better for handwriting and complex layouts")
    else:
        st.sidebar.warning("⚠️ **Note:** Tesseract works best with printed/typed text. For handwriting, try PaddleOCR or EasyOCR.")
    
    # File uploader
    uploaded_file = st.file_uploader(
        "Upload an image file (JPG/PNG)",
        type=['jpg', 'jpeg', 'png'],
        help="Upload a single-page English document image"
    )
    
    if uploaded_file is not None:
        # Display uploaded image
        try:
            image = Image.open(uploaded_file)
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("📷 Uploaded Image")
                st.image(image, caption="Original Image", use_container_width=True)
            
            with col2:
                st.subheader("⚙️ Processing Options")
                st.info("""
                **Image Info:**
                - Format: {}
                - Size: {} x {} pixels
                - Mode: {}
                """.format(
                    image.format,
                    image.width,
                    image.height,
                    image.mode
                ))
            
            # Convert button
            if st.button("🔄 Convert to Word Document", type="primary", use_container_width=True):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                try:
                    status_text.text("Preprocessing image...")
                    progress_bar.progress(10)
                    
                    # Preprocess image if enabled
                    if use_preprocessing:
                        processed_image = ocr_engine.preprocess_image(image, method=preprocessing_method)
                    else:
                        processed_image = image
                    
                    progress_bar.progress(30)
                    status_text.text("Running OCR...")
                    
                    # Extract text with or without formatting
                    doc = None
                    
                    # Choose OCR engine
                    if use_paddleocr:
                        # Use PaddleOCR (best overall accuracy)
                        # PaddleOCR often works better with original image, not preprocessed
                        ocr_image = image if not use_preprocessing else processed_image
                        
                        if use_formatting:
                            try:
                                text_blocks = ocr_engine.extract_text_with_formatting_paddleocr(ocr_image, min_confidence=min_confidence)
                            except Exception as e:
                                st.error(f"PaddleOCR error: {str(e)}")
                                st.info("Trying simple extraction instead...")
                                text_blocks = []
                            if text_blocks:
                                doc = create_docx_from_blocks(text_blocks)
                                st.subheader("📝 Extracted Text Preview")
                                preview_text = "\n".join([block.text for block in text_blocks])
                                st.text_area("Preview", preview_text, height=200, disabled=True)
                                
                                # Formatting summary
                                bold_count = sum(1 for block in text_blocks if block.is_bold)
                                italic_count = sum(1 for block in text_blocks if block.is_italic)
                                
                                col1, col2, col3 = st.columns(3)
                                with col1:
                                    st.metric("Total Blocks", len(text_blocks))
                                with col2:
                                    st.metric("Bold Detected", bold_count)
                                with col3:
                                    st.metric("Italic Detected", italic_count)
                            else:
                                # Fallback to simple PaddleOCR extraction
                                st.warning("No text detected with formatting. Trying simple extraction...")
                                try:
                                    text = ocr_engine.extract_text_paddleocr(ocr_image, min_confidence=min_confidence)
                                except Exception as e:
                                    st.error(f"PaddleOCR extraction failed: {str(e)}")
                                    text = ""
                                if text:
                                    doc = create_docx_simple(text)
                                    st.subheader("📝 Extracted Text")
                                    st.text_area("Extracted Text", text, height=200, disabled=True)
                                else:
                                    st.error("Could not extract any text from the image.")
                                    st.stop()
                        else:
                            # Simple PaddleOCR extraction
                            try:
                                text = ocr_engine.extract_text_paddleocr(ocr_image, min_confidence=min_confidence)
                            except Exception as e:
                                st.error(f"PaddleOCR extraction failed: {str(e)}")
                                text = ""
                            if text:
                                doc = create_docx_simple(text)
                                st.subheader("📝 Extracted Text")
                                st.text_area("Preview", text, height=200, disabled=True)
                            else:
                                st.error("Could not extract any text from the image.")
                                st.stop()
                    elif use_easyocr:
                        # Use EasyOCR (better for handwriting)
                        # EasyOCR often works better with original image, not preprocessed
                        ocr_image = image if not use_preprocessing else processed_image
                        
                        if use_formatting:
                            try:
                                text_blocks = ocr_engine.extract_text_with_formatting_easyocr(ocr_image, min_confidence=min_confidence)
                            except Exception as e:
                                st.error(f"EasyOCR error: {str(e)}")
                                st.info("Trying simple extraction instead...")
                                text_blocks = []
                            if text_blocks:
                                doc = create_docx_from_blocks(text_blocks)
                                st.subheader("📝 Extracted Text Preview")
                                preview_text = "\n".join([block.text for block in text_blocks])
                                st.text_area("Preview", preview_text, height=200, disabled=True)
                                
                                # Formatting summary
                                bold_count = sum(1 for block in text_blocks if block.is_bold)
                                italic_count = sum(1 for block in text_blocks if block.is_italic)
                                
                                col1, col2, col3 = st.columns(3)
                                with col1:
                                    st.metric("Total Blocks", len(text_blocks))
                                with col2:
                                    st.metric("Bold Detected", bold_count)
                                with col3:
                                    st.metric("Italic Detected", italic_count)
                            else:
                                # Fallback to simple EasyOCR extraction
                                st.warning("No text detected with formatting. Trying simple extraction...")
                                try:
                                    text = ocr_engine.extract_text_easyocr(ocr_image, min_confidence=min_confidence)
                                except Exception as e:
                                    st.error(f"EasyOCR extraction failed: {str(e)}")
                                    text = ""
                                if text:
                                    doc = create_docx_simple(text)
                                    st.subheader("📝 Extracted Text")
                                    st.text_area("Extracted Text", text, height=200, disabled=True)
                                else:
                                    st.error("Could not extract any text from the image.")
                                    st.stop()
                        else:
                            # Simple EasyOCR extraction
                            try:
                                text = ocr_engine.extract_text_easyocr(ocr_image, min_confidence=min_confidence)
                            except Exception as e:
                                st.error(f"EasyOCR extraction failed: {str(e)}")
                                text = ""
                            if text:
                                doc = create_docx_simple(text)
                                st.subheader("📝 Extracted Text")
                                st.text_area("Preview", text, height=200, disabled=True)
                            else:
                                st.error("Could not extract any text from the image.")
                                st.stop()
                    else:
                        # Use Tesseract OCR
                        # Note: When try_multiple_modes is enabled, we skip formatting detection for simplicity
                        if use_formatting and not try_multiple_modes:
                            # Use formatting detection with single PSM mode
                            text_blocks = ocr_engine.extract_text_with_formatting(processed_image, psm_mode=psm_mode)
                            
                            if text_blocks:
                                # Create formatted document
                                doc = create_docx_from_blocks(text_blocks)
                                
                                # Display extracted text preview
                                st.subheader("📝 Extracted Text Preview")
                                preview_text = "\n".join([block.text for block in text_blocks])
                                st.text_area(
                                    "Preview",
                                    preview_text,
                                    height=200,
                                    disabled=True
                                )
                                
                                # Formatting summary
                                bold_count = sum(1 for block in text_blocks if block.is_bold)
                                italic_count = sum(1 for block in text_blocks if block.is_italic)
                                
                                col1, col2, col3 = st.columns(3)
                                with col1:
                                    st.metric("Total Blocks", len(text_blocks))
                                with col2:
                                    st.metric("Bold Detected", bold_count)
                                with col3:
                                    st.metric("Italic Detected", italic_count)
                            else:
                                # Fallback to simple extraction
                                st.warning("No text detected with formatting. Trying simple extraction...")
                                if try_multiple_modes:
                                    text, best_psm = ocr_engine.try_multiple_psm_modes(processed_image)
                                    st.info(f"✅ Best result found using PSM mode {best_psm}")
                                else:
                                    text = ocr_engine.extract_text_simple(processed_image, psm_mode=psm_mode)
                                if text:
                                    doc = create_docx_simple(text)
                                    st.subheader("📝 Extracted Text")
                                    st.text_area("Extracted Text", text, height=200, disabled=True)
                                else:
                                    st.error("Could not extract any text from the image.")
                                    st.stop()
                        else:
                            # Simple extraction without formatting
                            if try_multiple_modes:
                                text, best_psm = ocr_engine.try_multiple_psm_modes(processed_image)
                                st.info(f"✅ Best result found using PSM mode {best_psm}")
                            else:
                                text = ocr_engine.extract_text_simple(processed_image, psm_mode=psm_mode)
                            if text:
                                doc = create_docx_simple(text)
                                st.subheader("📝 Extracted Text")
                                st.text_area("Preview", text, height=200, disabled=True)
                            else:
                                st.error("Could not extract any text from the image.")
                                st.stop()
                        
                        # Save document to bytes (only if doc was created)
                        progress_bar.progress(90)
                        status_text.text("Creating document...")
                        
                        if doc is None:
                            st.error("Document creation failed.")
                            progress_bar.empty()
                            status_text.empty()
                            st.stop()
                        
                        doc_bytes = io.BytesIO()
                        doc.save(doc_bytes)
                        doc_bytes.seek(0)
                        
                        progress_bar.progress(100)
                        progress_bar.empty()
                        status_text.empty()
                        
                        # Download button
                        st.success("✅ Conversion successful!")
                        st.download_button(
                            label="📥 Download Word Document (.docx)",
                            data=doc_bytes,
                            file_name="converted_document.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                        
                except Exception as e:
                    progress_bar.empty()
                    status_text.empty()
                    st.error(f"❌ Error during processing: {str(e)}")
                    with st.expander("🔍 Error Details"):
                        st.exception(e)
                    st.info("💡 **Tips to improve results:**\n"
                           "- Try disabling preprocessing for EasyOCR\n"
                           "- Adjust confidence threshold (lower = more text, higher = more accurate)\n"
                           "- Ensure image is clear and well-lit\n"
                           "- Try Tesseract for printed text\n"
                           "- For handwriting, EasyOCR works best with original images")
        
        except Exception as e:
            st.error(f"Error loading image: {str(e)}")
            st.exception(e)
    
    else:
        # Show instructions when no file is uploaded
        st.info("👆 Please upload an image file to get started.")
        
        st.markdown("""
        ### How to use:
        1. **Upload** a JPG or PNG image containing text
        2. **Preview** the uploaded image
        3. **Adjust settings** in the sidebar (optional)
        4. **Click** the "Convert to Word Document" button
        5. **Download** the resulting .docx file
        
        ### Features:
        - ✅ OCR text extraction using Tesseract
        - ✅ Formatting detection (Bold, Italic)
        - ✅ Alignment preservation (Left, Center, Right)
        - ✅ Image preprocessing for better accuracy
        - ✅ Clean, user-friendly interface
        
        ### Tips for best results:
        - Use clear, high-resolution images
        - Ensure good contrast between text and background
        - Use single-page English documents
        - Enable preprocessing for scanned documents
        """)


if __name__ == "__main__":
    main()
